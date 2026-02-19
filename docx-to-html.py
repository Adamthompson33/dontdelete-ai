from docx import Document
import html

doc = Document(r"C:\Users\adamt\Downloads\substrate-project-summary.docx")

lines = []
lines.append("""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>The Substrate — Project Summary</title>
<style>
  body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 800px; margin: 0 auto; padding: 40px 20px; line-height: 1.6; color: #1a1a2e; background: #fafafa; }
  h1 { font-size: 28px; color: #1a1a2e; border-bottom: 3px solid #00d4aa; padding-bottom: 10px; }
  h2 { font-size: 22px; color: #1a1a2e; margin-top: 30px; border-bottom: 1px solid #ddd; padding-bottom: 6px; }
  h3 { font-size: 18px; color: #333; margin-top: 20px; }
  p { margin: 10px 0; }
  table { border-collapse: collapse; width: 100%; margin: 15px 0; }
  th, td { border: 1px solid #ddd; padding: 8px 12px; text-align: left; }
  th { background: #1a1a2e; color: white; }
  tr:nth-child(even) { background: #f5f5f5; }
  ul, ol { margin: 10px 0; padding-left: 25px; }
  li { margin: 4px 0; }
  .footer { margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; color: #666; font-size: 14px; }
</style>
</head>
<body>
""")

for para in doc.paragraphs:
    text = html.escape(para.text.strip())
    if not text:
        continue
    
    style = para.style.name if para.style else ""
    
    if "Heading 1" in style:
        lines.append(f"<h1>{text}</h1>")
    elif "Heading 2" in style:
        lines.append(f"<h2>{text}</h2>")
    elif "Heading 3" in style:
        lines.append(f"<h3>{text}</h3>")
    elif "List" in style or text.startswith("•") or text.startswith("-"):
        lines.append(f"<li>{text.lstrip('•- ')}</li>")
    else:
        # Bold detection
        bold_parts = []
        for run in para.runs:
            t = html.escape(run.text)
            if run.bold:
                bold_parts.append(f"<strong>{t}</strong>")
            else:
                bold_parts.append(t)
        if bold_parts:
            lines.append(f"<p>{''.join(bold_parts)}</p>")
        else:
            lines.append(f"<p>{text}</p>")

# Handle tables
for table in doc.tables:
    lines.append("<table>")
    for i, row in enumerate(table.rows):
        lines.append("<tr>")
        tag = "th" if i == 0 else "td"
        for cell in row.cells:
            lines.append(f"<{tag}>{html.escape(cell.text)}</{tag}>")
        lines.append("</tr>")
    lines.append("</table>")

lines.append("""
<div class="footer">
  <p>The Substrate — Project Summary | February 2026</p>
</div>
</body>
</html>""")

output = r"C:\Users\adamt\Downloads\substrate-project-summary.html"
with open(output, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))

print(f"Done! Saved to {output}")
